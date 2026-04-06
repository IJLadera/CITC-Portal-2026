from bdb import effective
from rest_framework import viewsets, status
from rest_framework.exceptions import PermissionDenied
from rest_framework.decorators import action
from rest_framework.response import Response    
from django.contrib.contenttypes.models import ContentType
from django.db.models import F, OuterRef, Subquery, Q, Count
from django.db import transaction
from django.utils import timezone
from django.http import FileResponse, JsonResponse
from django.conf import settings 
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
from urllib.parse import urlsplit, urlunsplit
import tempfile
import requests
import io

from .models import  ( 
    Syllabus, 
    SyllabusInstructor, 
    SyllabusCourseOutcome, 
    SyllCoPo, 
    SyllabusCourseOutline,  
    SyllabusDeanFeedback, 
    SyllabusCotCo, 
    SyllabusTemplate,  
    SyllabusComment,
    ReviewFormTemplate,
    ReviewFormField,
    ReviewFormItem,
    SRFForm,
    SRFIndicator,
    SRFFieldValue
)

from .serializers import (
    SRFFormSerializer,
    SyllabusCreateSerializer,
    SyllabusUpdateSerializer,
    SyllabusVersionSerializer,
    SyllabusListSerializer,
    SyllabusDetailSerializer,
    SyllabusCourseOutcomeSerializer,
    SyllCoPoSerializer,
    SyllabusCourseOutlineSerializer,
    SyllabusCourseRequirementSerializer, 
    SyllabusApprovedReadSerializer,
    SyllabusTemplateSerializer,
    SyllabusCommentSerializer,
    ReviewFormTemplateSerializer,
    SRFIndicatorSerializer, 
    get_current_chair_json,
    get_current_dean_json
)

from .pagination import SyllabiPagination
from .utils.prefill_utils import get_prefill_value
from auditlog.models import LogEntry
from academics.models import PEO, ProgramOutcome 
from bayanihan.models import BayanihanGroupUser, BayanihanGroup 
from users.models import Role, UserRole
from shared.models import Report

from users.permissions import RolePermission

import os
import re
from datetime import date
from docxtpl import DocxTemplate, RichText, InlineImage, Subdoc  
from docx.shared import Pt, Mm, Inches, RGBColor 
from docx.enum.text import WD_ALIGN_PARAGRAPH 
from docx.oxml.ns import qn, nsdecls 
from docx.enum.table import WD_ALIGN_VERTICAL
from urllib.parse import urlparse  
from docx.oxml import parse_xml     
import html as ihtml 
from bs4.element import NavigableString 
import html as ihtml
from bs4 import BeautifulSoup 

from tos.views import convert_docx_to_pdf
import sys

if sys.platform == "win32":
    import pythoncom
    from docx2pdf import convert
else:
    pythoncom = None
    convert = None
    
import platform
import subprocess

def get_public_url(url: str):
    parts = urlsplit(url)
    # Remove querystring (signed params)
    clean = urlunsplit((parts.scheme, parts.netloc, parts.path, "", ""))
    return clean

# Create your views here. 
class SyllabusViewSet(viewsets.ModelViewSet):
    permission_classes = [RolePermission(
        "ADMIN", "BAYANIHAN_LEADER", "BAYANIHAN_TEACHER", 
        "DEAN", "CHAIRPERSON", "AUDITOR"
    )]
    pagination_class = SyllabiPagination

    def get_queryset(self):
        user = self.request.user
        role = self.request.GET.get("role")

        # Base queryset
        qs = Syllabus.objects.select_related(
            "syllabus_template", "bayanihan_group", "course", "college", "program__department", "curriculum"
        ).prefetch_related(
            "peos", "program_outcomes", "instructors__user", "course_outcomes",
            "syllcopos", "course_outlines", "dean_feedback", "review_form"
        )

        # Enforce role filtering only for list
        if self.action in ["list"]:
            if not role:
                raise PermissionDenied("Missing role parameter.")
            role.upper()

            # --- ADMIN ---
            if role == "ADMIN":
                if not user.has_role("ADMIN"):
                    raise PermissionDenied("You are not an admin.")

                latest_sub = (
                    Syllabus.objects.filter(bayanihan_group_id=OuterRef("bayanihan_group_id"))
                    .order_by("-version")
                    .values("version")[:1]
                )

                qs = qs.annotate(latest_version=Subquery(latest_sub)).filter(
                    version=F("latest_version")
                )   

            # --- BAYANIHAN LEADER ---
            elif role == "BAYANIHAN_LEADER":
                leader_groups = BayanihanGroupUser.objects.filter(
                    user=user, role="LEADER"
                ).values_list("group_id", flat=True)

                if not leader_groups:
                    raise PermissionDenied("You are not a leader in any Bayanihan group.")

                latest_sub = (
                    Syllabus.objects.filter(bayanihan_group_id=OuterRef("bayanihan_group_id"))
                    .order_by("-version")
                    .values("version")[:1]
                )

                qs = qs.annotate(latest_version=Subquery(latest_sub)).filter(
                    bayanihan_group_id__in=leader_groups,
                    version=F("latest_version")
                )

            # --- BAYANIHAN TEACHER ---
            elif role == "BAYANIHAN_TEACHER":
                teacher_groups = BayanihanGroupUser.objects.filter(
                    user=user, role="TEACHER"
                ).values_list("group_id", flat=True)

                if not teacher_groups:
                    raise PermissionDenied("You are not a teacher in any Bayanihan group.")

                latest_sub = (
                    Syllabus.objects.filter(bayanihan_group_id=OuterRef("bayanihan_group_id"))
                    .order_by("-version")
                    .values("version")[:1]
                )

                qs = qs.annotate(latest_version=Subquery(latest_sub)).filter(
                    bayanihan_group_id__in=teacher_groups,
                    version=F("latest_version")
                )

            # --- CHAIRPERSON ---
            elif role == "CHAIRPERSON":
                try:
                    chair_role = UserRole.objects.get(
                        user=user,
                        entity_type="Department",
                        role__name="CHAIRPERSON",
                        entity_id__isnull=False,
                    )
                except UserRole.DoesNotExist:
                    raise PermissionDenied("You are not a Chairperson.")
                except UserRole.MultipleObjectsReturned:
                    raise PermissionDenied("Multiple Chairperson roles found. Contact admin.")

                latest_sub = (
                    Syllabus.objects.filter(
                        bayanihan_group_id=OuterRef("bayanihan_group_id"),
                        chair_submitted_at__isnull=False,
                    )
                    .order_by("-version")
                    .values("version")[:1]
                )

                qs = qs.annotate(latest_chair_version=Subquery(latest_sub)).filter(
                    program__department_id=chair_role.entity_id,
                    version=F("latest_chair_version")
                )

            # --- DEAN ---
            elif role == "DEAN":
                try:
                    dean_role = UserRole.objects.get(
                        user=user,
                        entity_type="College",
                        role__name="DEAN",
                        entity_id__isnull=False,
                    )
                except UserRole.DoesNotExist:
                    raise PermissionDenied("You are not a Dean.")
                except UserRole.MultipleObjectsReturned:
                    raise PermissionDenied("Multiple Dean roles found. Contact admin.")

                latest_sub = (
                    Syllabus.objects.filter(
                        bayanihan_group_id=OuterRef("bayanihan_group_id"),
                        dean_submitted_at__isnull=False,
                    )
                    .order_by("-version")
                    .values("version")[:1]
                )

                qs = qs.annotate(latest_dean_version=Subquery(latest_sub)).filter(
                    college_id=dean_role.entity_id,
                    version=F("latest_dean_version")
                )

            # --- AUDITOR ---
            elif role == "AUDITOR":
                if not UserRole.objects.filter(user=user, role__name="AUDITOR").exists():
                    raise PermissionDenied("You are not an Auditor.")

                latest_sub = (
                    Syllabus.objects.filter(
                        bayanihan_group_id=OuterRef("bayanihan_group_id"),
                        dean_approved_at__isnull=False,
                    )
                    .order_by("-version")
                    .values("version")[:1]
                )

                qs = qs.annotate(latest_version=Subquery(latest_sub)).filter(
                    version=F("latest_version")
                )

            else:
                raise PermissionDenied("Invalid role parameter.")

        # === Apply shared filters ===
        year_level = self.request.GET.get("year_level")
        department = self.request.GET.get("department")
        program = self.request.GET.get("program")
        semester = self.request.GET.get("semester")
        school_year = self.request.GET.get("school_year")
        status = self.request.GET.get("status")
        search = self.request.GET.get("search")

        if year_level:
            qs = qs.filter(course__course_year_level__iexact=year_level) 
        if department:
            qs = qs.filter(program__department_id=department) 
        if program:
            qs = qs.filter(program_id=program)
        if semester:
            qs = qs.filter(course__course_semester__iexact=semester)
        if school_year:
            qs = qs.filter(bayanihan_group__school_year__iexact=school_year)
        if status:
            qs = qs.filter(status__iexact=status)
        if search:
            qs = qs.filter(
                Q(course__course_code__icontains=search)
                | Q(course__course_title__icontains=search)
            )

        # === Optional: return all if ?all=true (for dropdowns) ===
        if self.request.GET.get("all") == "true":
            self.pagination_class = None

        return qs.order_by("-updated_at")

    def get_serializer_class(self):
        if self.action == "create":
            return SyllabusCreateSerializer
        if self.action in ["update", "partial_update"]:
            return SyllabusUpdateSerializer
        if self.action == "list":
            return SyllabusListSerializer
        if self.action == "retrieve":
            return SyllabusDetailSerializer 
        return SyllabusDetailSerializer 

    @action(detail=True, methods=["get"], url_path="audit-logs")
    def get_audit_logs(self, request, pk=None):
        syllabus = self.get_object()

        logs = LogEntry.objects.get_for_object(syllabus).select_related("actor").order_by("-timestamp")

        results = []
        for log in logs:
            results.append({
                "id": log.id,
                "event": log.get_action_display(),  # CREATE / UPDATE / DELETE
                "user": {
                    "firstname": getattr(log.actor, "first_name", ""),
                    "lastname": getattr(log.actor, "last_name", ""),
                },
                "changes": {
                    field: {"old": values[0], "new": values[1]}
                    for field, values in log.changes_dict.items()
                },
                "created_at": log.timestamp,
            })

        return Response(results, status=status.HTTP_200_OK)
        
    @action(detail=True, methods=["get"], url_path="syllabus-versions")
    def get_syllabus_versions(self, request, pk=None):
        syllabus = self.get_object()

        try: 
            syllabus_versions = Syllabus.objects.filter(
                bayanihan_group_id=syllabus.bayanihan_group_id
            ).order_by("-version")  # optional: order by version desc
        except Syllabus.DoesNotExist:
            return Response(
                {"detail": "No versions found for this syllabus."},
                status=status.HTTP_404_NOT_FOUND,
            )

        serializer = SyllabusVersionSerializer(syllabus_versions, many=True)
        return Response(serializer.data, status=status.HTTP_200_OK)
    
    @action(detail=True, methods=["patch"], url_path="course-requirements")
    def update_course_requirements(self, request, pk=None):
        syllabus = self.get_object()
        user = request.user
        
        role = request.query_params.get("role")      
        if not role:
            raise PermissionDenied("Role parameter is required.")
        role = role.upper()  # normalize 
        allowed = False 
        if role == "ADMIN" and UserRole.objects.filter(user=user, role__name="ADMIN").exists():
            allowed = True
        # BAYANIHAN_LEADER can edit course_requirements only for syllabus that belong to their groups
        elif role == "BAYANIHAN_LEADER" and UserRole.objects.filter(user=user, role__name="BAYANIHAN_LEADER").exists():
            leader_group_ids = BayanihanGroupUser.objects.filter(
                user=user, role="LEADER"
            ).values_list("group_id", flat=True)
            if syllabus.bayanihan_group.id in leader_group_ids:
                allowed = True

        if not allowed:
            raise PermissionDenied("You do not have permission to edit this syllabus' course requirements.")
        
        serializer = SyllabusCourseRequirementSerializer(
            syllabus, data=request.data, partial=True
        )
        serializer.is_valid(raise_exception=True)
        serializer.save()
        return Response(serializer.data, status=status.HTTP_200_OK)
    
    @action(detail=True, methods=["patch"], url_path="submit-syllabus")
    @transaction.atomic
    def submit_syllabus(self, request, pk=None):
        syllabus = self.get_object()
        user = request.user
        role = request.query_params.get("role")     

        if not role:
            raise PermissionDenied("Role parameter is required.")
        role = role.upper()  # normalize
        allowed = False
        if role == "ADMIN" and UserRole.objects.filter(user=user, role__name="ADMIN").exists():
            allowed = True
        # BAYANIHAN_LEADER can submit only for their groups
        elif role == "BAYANIHAN_LEADER" and UserRole.objects.filter(user=user, role__name="BAYANIHAN_LEADER").exists():
            leader_group_ids = BayanihanGroupUser.objects.filter(
                user=user, role="LEADER"
            ).values_list("group_id", flat=True)
            if syllabus.bayanihan_group.id in leader_group_ids:
                allowed = True

        if not allowed:
            raise PermissionDenied("You do not have permission to submit this syllabus.")

        # Status update logic
        if syllabus.status == "Draft":
            syllabus.status = "Pending Chair Review"
            syllabus.chair_submitted_at = timezone.now()
        elif syllabus.status == "Requires Revision":
            syllabus.status = "Revisions Applied"
            syllabus.chair_submitted_at = timezone.now()
        else:
            return Response(
                {"detail": "Syllabus cannot be submitted from its current status."},
                status=status.HTTP_400_BAD_REQUEST
            ) 
        syllabus.save()

        # --- ✅ Update or create Report record for this Syllabus ---
        try:
            report = Report.objects.get(syllabus=syllabus)
        except Report.DoesNotExist:
            report = Report.objects.create(
                bayanihan_group=syllabus.bayanihan_group,
                syllabus=syllabus,
                version=syllabus.version,
            )  
        report.mark_chair_submitted()

        # Return minimal info via serializer
        serializer = SyllabusDetailSerializer(syllabus, context={"request": request})
        return Response(serializer.data, status=status.HTTP_200_OK)
    
    @action(detail=True, methods=["post"], url_path="review")
    @transaction.atomic
    def review_syllabus(self, request, pk=None):
        """
        Unified endpoint for Chairperson to review a syllabus.
        Combines logic from the old review-syllabus-chair and submit_review actions.

        Expected payload:
        {
          "action": 1,  # 1 = approve, 0 = reject
          "checklist": [
            {"item": 10, "response": "yes", "remarks": ""},
            {"item": 11, "response": "no", "remarks": "Needs fix"}
          ],
          "fields": [
            {"field": 3, "value": "CPE101"},
            {"field": 4, "value": "Computer Programming 1"}
          ]
        }
        """
        syllabus = self.get_object()
        user = request.user
        role = request.query_params.get("role")

        if not role:
            raise PermissionDenied("Role parameter is required.")

        department = syllabus.program.department 
        # ✅ Role-based permission check
        allowed = False
        if role == "ADMIN" and UserRole.objects.filter(user=user, role__name="ADMIN").exists():
            allowed = True
        elif role == "CHAIRPERSON":
            chair_dept_ids = UserRole.objects.filter(
                user=user,
                entity_type="Department",
                role__name="CHAIRPERSON",
            ).values_list("entity_id", flat=True)
            if department.id in chair_dept_ids:
                allowed = True

        if not allowed:
            raise PermissionDenied("You do not have permission to review this syllabus.")
        
        # Get Department Chairperson
        chair_role = UserRole.objects.filter(
            entity_type="Department",
            entity_id=department.id,
            role__name="CHAIRPERSON",
        ).select_related("user").first()

        if not chair_role:
            raise PermissionDenied("No Chairperson is assigned to this Department.")

        chair_user = chair_role.user
        user = chair_user

        data = request.data
        action_value = data.get("action")
        checklist_data = data.get("checklist", [])
        field_data = data.get("fields", [])

        # ✅ Validate action
        if action_value not in [0, 1]:
            return Response({"detail": "Invalid action. Must be 0 (reject) or 1 (approve)."},
                            status=status.HTTP_400_BAD_REQUEST)

        # ✅ Get latest active review form template
        form_template = ReviewFormTemplate.objects.filter(is_active=True).order_by("-revision_no").first()
        if not form_template:
            return Response({"detail": "No active review form template found. Please tell Admin to create a template."},
                            status=status.HTTP_400_BAD_REQUEST)

        # ✅ Update Syllabus status and timestamps
        if action_value == 1:
            syllabus.status = "Approved by Chair"
            syllabus.dean_submitted_at = timezone.now()
        else:
            syllabus.status = "Returned by Chair"
            syllabus.chair_rejected_at = timezone.now()
        syllabus.save()

        # ✅ Reviewer info snapshot
        reviewed_by = (
            f"{user.prefix or ''} {user.first_name} {user.last_name} {user.suffix or ''}"
        ).strip() or user.username 

        # ✅ Create SRFForm (main review form record)
        srf_form = SRFForm.objects.create(
            form_template=form_template,
            effective_date=form_template.effective_date if form_template.effective_date else None,
            syllabus=syllabus,
            user=user,
            action=action_value,
            reviewed_by_snapshot=reviewed_by, 
        )

        # ✅ Create SRFIndicators (checklist responses)
        for item in checklist_data:
            SRFIndicator.objects.create(
                review_form=srf_form,
                item_id=item["item"],
                response=item.get("response"),
                remarks=item.get("remarks", ""),
            )

        # --- Create SRFFieldValues (prefill + frontend) ---
        for field in form_template.fields.all():
            # Check if frontend submitted a value
            frontend_field = next((f for f in field_data if f["field"] == field.id), None)

            if field.prefill_source != "none":
                # Use prefill value from syllabus
                value = get_prefill_value(field, syllabus)
            elif frontend_field:
                # Use value submitted by frontend
                value = frontend_field.get("value", "")
            else:
                # Default empty
                value = ""

            SRFFieldValue.objects.create(
                review_form=srf_form,
                field_id=field.id,
                value=value,
            ) 

        # ✅ Ensure Report record exists
        report, _ = Report.objects.get_or_create(
            syllabus=syllabus,
            defaults={
                "bayanihan_group": syllabus.bayanihan_group,
                "version": syllabus.version,
            },
        ) 
        if action_value == 1:
            report.mark_dean_submitted()
        else:
            report.mark_chair_rejected()

        return Response(
            {
                "detail": f"Syllabus {'approved' if action_value == 1 else 'rejected'} successfully.",
                "srf_form_id": srf_form.id,
                "syllabus_id": syllabus.id,
            },
            status=status.HTTP_201_CREATED,
        )
        
    @action(detail=True, methods=["patch"], url_path="review-syllabus-dean")
    @transaction.atomic
    def review_syllabus_dean(self, request, pk=None):
        syllabus = self.get_object()
        user = request.user  
        
        if UserRole.objects.filter(user=user, role__name="ADMIN").exists():
            allowed = True
        elif UserRole.objects.filter(user=user, role__name="DEAN").exists():
            dean_college_ids = UserRole.objects.filter(
                user=user,
                entity_type="College",
                role__name="DEAN",
            ).values_list("entity_id", flat=True)
            if syllabus.college_id in dean_college_ids:
                allowed = True

        if not allowed:
            raise PermissionDenied("You do not have permission to review this syllabus as Dean.")

        decision = request.data.get("decision")
        if decision not in ["approve", "reject"]:
            return Response(
                {"detail": "Invalid decision. Must be 'approve' or 'reject'."},
                status=status.HTTP_400_BAD_REQUEST,
            )

        if decision == "approve":
            syllabus.status = "Approved by Dean"
            syllabus.dean_approved_at = timezone.now()
            syllabus.save()

        else:  # reject
            feedback_text = request.data.get("feedback_text")
            if not feedback_text:
                return Response(
                    {"detail": "Feedback text is required when rejecting."},
                    status=status.HTTP_400_BAD_REQUEST,
                )

            syllabus.status = "Returned by Dean"
            syllabus.dean_rejected_at = timezone.now()
            syllabus.save()

            # create dean feedback record
            SyllabusDeanFeedback.objects.create(
                syllabus=syllabus,
                user=user,
                feedback_text=feedback_text,
            )  

        # --- ✅ Update corresponding Report record --- 
        try:
            report = Report.objects.get(syllabus=syllabus)
        except Report.DoesNotExist:
            report = Report.objects.create(
                bayanihan_group=syllabus.bayanihan_group,
                syllabus=syllabus,
                version=syllabus.version,
            ) 
        if decision == "approve":
            report.mark_dean_approved()   # ⬅ convenience method
        else:
            report.mark_dean_rejected() 
                
        syllabus_data = SyllabusDetailSerializer(syllabus).data
        return Response(syllabus_data, status=status.HTTP_201_CREATED)
        
    @action(detail=True, methods=["get"], url_path="review-form")
    def get_review_form(self, request, pk=None):
        syllabus = self.get_object()

        try:
            review_form = SRFForm.objects.select_related("syllabus", "user").prefetch_related("indicators", "field_values").get(
                syllabus=syllabus
            )
        except SRFForm.DoesNotExist:
            return Response(
                {"detail": "No review form found for this syllabus."},
                status=status.HTTP_404_NOT_FOUND,
            )

        serializer = SRFFormSerializer(review_form)
        return Response(serializer.data, status=status.HTTP_200_OK)
    
    @action(detail=True, methods=["post"], url_path="replicate-syllabus")
    @transaction.atomic
    def replicate_syllabus(self, request, pk=None):
        syllabus = self.get_object()
        user = request.user
 
        if syllabus.status not in ["Returned by Chair", "Returned by Dean"]:
            return Response(
                {"detail": "Replication allowed only for returned syllabi."},
                status=status.HTTP_400_BAD_REQUEST,
            )

        # ✅ Only Bayanihan Leader of that group (or Admin) can replicate
        role = request.query_params.get("role", "").upper()
        allowed = False
        if role == "ADMIN" and UserRole.objects.filter(user=user, role__name="ADMIN").exists():
            allowed = True
        elif role == "BAYANIHAN_LEADER":
            leader_group_ids = BayanihanGroupUser.objects.filter(
                user=user, role="LEADER"
            ).values_list("group_id", flat=True)
            if syllabus.bayanihan_group_id in leader_group_ids:
                allowed = True

        if not allowed:
            raise PermissionDenied("You do not have permission to replicate this syllabus.")
 
        new_syllabus = Syllabus.objects.create(
            syllabus_template=syllabus.syllabus_template,
            effective_date=syllabus.effective_date,
            bayanihan_group=syllabus.bayanihan_group,
            course=syllabus.course,
            college=syllabus.college,
            program=syllabus.program,
            curriculum=syllabus.curriculum,
            
            class_schedules=syllabus.class_schedules,
            building_room=syllabus.building_room,
            class_contact=syllabus.class_contact,
            
            consultation_hours=syllabus.consultation_hours,
            consultation_room=syllabus.consultation_room,
            consultation_contact=syllabus.consultation_contact,
            
            course_description=syllabus.course_description,
            
            course_requirements=syllabus.course_requirements,
            
            version=syllabus.version + 1,
            status="Requires Revision",
            
            chair_submitted_at=None,
            chair_rejected_at=None,
            dean_submitted_at=None,
            dean_rejected_at=None,
            dean_approved_at=None,
            
            dean=syllabus.dean,
            chair=syllabus.chair
        )

        # ✅ Clone ManyToMany fields (PEOs and Program Outcomes)
        new_syllabus.peos.set(syllabus.peos.all())
        new_syllabus.program_outcomes.set(syllabus.program_outcomes.all())

        # ✅ Clone related instructors
        instructors = SyllabusInstructor.objects.filter(syllabus=syllabus)
        for inst in instructors:
            SyllabusInstructor.objects.create(
                syllabus=new_syllabus,
                user=inst.user, 
            )

        # ✅ Clone course outcomes
        outcomes = SyllabusCourseOutcome.objects.filter(syllabus=syllabus)
        for outcome in outcomes:
            new_outcome = SyllabusCourseOutcome.objects.create(
                syllabus=new_syllabus,
                co_code=outcome.co_code,
                co_description=outcome.co_description,
            )

            # ✅ Also copy linked CoPo mappings
            copos = SyllCoPo.objects.filter(course_outcome=outcome)
            for copo in copos:
                SyllCoPo.objects.create(
                    syllabus=new_syllabus,
                    course_outcome=new_outcome,
                    program_outcome=copo.program_outcome,
                    syllabus_co_po_code=copo.syllabus_co_po_code,
                )

        # ✅ Clone course outlines
        outlines = SyllabusCourseOutline.objects.filter(syllabus=syllabus)
        for outline in outlines:
            new_outline = SyllabusCourseOutline.objects.create(
                syllabus=new_syllabus,
                syllabus_term=outline.syllabus_term,
                row_no=outline.row_no, 
                allotted_hour=outline.allotted_hour,
                allotted_time=outline.allotted_time,
                intended_learning=outline.intended_learning,
                topics=outline.topics,
                suggested_readings=outline.suggested_readings,
                learning_activities=outline.learning_activities,
                assessment_tools=outline.assessment_tools,
                grading_criteria=outline.grading_criteria,
                remarks=outline.remarks,
            )

            # ✅ Copy outline–CoPo mappings
            cotcos = SyllabusCotCo.objects.filter(course_outline=outline)
            for cotco in cotcos:
                SyllabusCotCo.objects.create( 
                    course_outline=new_outline,
                    course_outcome=cotco.course_outcome, 
                )

        # Create another Report record for Dean Reports (next version)
        Report.objects.create(
            syllabus=new_syllabus,
            bayanihan_group=new_syllabus.bayanihan_group, 
            version=new_syllabus.version,
        )
                
        new_syllabus_data = SyllabusDetailSerializer(new_syllabus).data
        return Response(new_syllabus_data, status=status.HTTP_201_CREATED)
    
    @action(detail=False, methods=["get"], url_path="past-syllabi/(?P<bayanihan_group_id>[^/.]+)")
    def past_syllabi(self, request, bayanihan_group_id=None): 
        user = request.user 
        role = request.query_params.get("role")     

        if not role:
            raise PermissionDenied("Role parameter is required.")
        role = role.upper() 
 
        try:
            bayanihan_group = BayanihanGroup.objects.select_related("course").get(id=bayanihan_group_id)
        except BayanihanGroup.DoesNotExist:
            return Response({"detail": "Bayanihan group not found."}, status=status.HTTP_404_NOT_FOUND)

        course_id = bayanihan_group.course.pk

        # 2️⃣ Filter based on role
        if role == "ADMIN":
            if not UserRole.objects.filter(user=user, role__name="ADMIN").exists():
                raise PermissionDenied("You are not an Admin.")

            qs = (
                Syllabus.objects
                .filter(
                    course_id=course_id,
                    status="Approved by Dean",
                    dean_approved_at__isnull=False,
                )
                .select_related("bayanihan_group", "course") 
                .order_by("-dean_approved_at")
            )

        elif role == "BAYANIHAN_LEADER":
            if not UserRole.objects.filter(user=user, role__name="BAYANIHAN_LEADER").exists():
                raise PermissionDenied("You are not a Bayanihan Leader.")

            # Verify user is a leader of this group
            is_leader = BayanihanGroupUser.objects.filter(
                user=user, group_id=bayanihan_group_id, role="LEADER"
            ).exists()

            if not is_leader:
                return Response(
                    {"detail": "You are not the leader of this Bayanihan group."},
                    status=403,
                )

            qs = (
                Syllabus.objects
                .filter(
                    course_id=course_id,
                    status="Approved by Dean",
                    dean_approved_at__isnull=False,
                )
                .select_related("bayanihan_group", "bayanihan_group__course") 
                .order_by("-dean_approved_at")
            )

        else:
            raise PermissionDenied("Invalid role parameter.")

        serializer = SyllabusListSerializer(qs, many=True)
        return Response(serializer.data, status=status.HTTP_200_OK)
   
    @action(detail=True, methods=["post"], url_path="duplicate-syllabus")
    @transaction.atomic
    def duplicate_syllabus(self, request, pk=None):
        syllabus = self.get_object()
        user = request.user
 
        if syllabus.status != "Approved by Dean" or syllabus.dean_approved_at is None:
            return Response(
                {"detail": "Duplicate allowed only for approved syllabi."},
                status=status.HTTP_400_BAD_REQUEST,
            )
 
        bayanihan_group_id = request.data.get("bayanihan_group")
        if not bayanihan_group_id:
            return Response(
                {"detail": "Please provide a valid Bayanihan Group."},
                status=status.HTTP_400_BAD_REQUEST,
            )

        try:
            target_group = BayanihanGroup.objects.select_related("course").get(id=bayanihan_group_id)
        except BayanihanGroup.DoesNotExist:
            return Response(
                {"detail": "Bayanihan Group doesn't exist."},
                status=status.HTTP_404_NOT_FOUND,
            )
            
        # ✅ Permission logic
        role = request.query_params.get("role", "").upper()
        allowed = False
        if role == "ADMIN" and UserRole.objects.filter(user=user, role__name="ADMIN").exists():
            allowed = True
        elif role == "BAYANIHAN_LEADER":
            leader_group_ids = BayanihanGroupUser.objects.filter(
                user=user, role="LEADER"
            ).values_list("group_id", flat=True)

            # ✅ BAYANIHAN_LEADER may only duplicate to a group with the same course.id
            if target_group.id in leader_group_ids and target_group.course.id == syllabus.course.id:
                allowed = True 

        if not allowed:
            raise PermissionDenied(
                "You do not have permission to duplicate this syllabus to the selected Bayanihan group."
            )

        # Take new effective date from current active Syllabus Template (e.g., December 30 of the end year)
        syllabus_template = (
            SyllabusTemplate.objects
            .filter(is_active=True)
            .order_by("-revision_no")
            .first()
        )

        # ✅ Rebuild dean/chair from current assignments
        request_obj = request  # for absolute URLs in signatures
        dean_json = get_current_dean_json(syllabus.college, request_obj)
        chair_json = get_current_chair_json(syllabus.program.department, request_obj)

        if dean_json is None or chair_json is None:
            missing = []
            if dean_json is None:
                missing.append("Dean")
            if chair_json is None:
                missing.append("Chairperson")
            raise ValidationError(
                f"Cannot duplicate syllabus: no {', '.join(missing)} assigned."
            )

        new_syllabus = Syllabus.objects.create(
            syllabus_template=syllabus_template,
            effective_date=syllabus_template.effective_date if syllabus_template and syllabus_template.effective_date else None,
            bayanihan_group=target_group,
            course=syllabus.course,
            college=syllabus.college,
            program=syllabus.program,
            curriculum=syllabus.curriculum,
            
            class_schedules=syllabus.class_schedules,
            building_room=syllabus.building_room,   

            consultation_hours=syllabus.consultation_hours,
            consultation_room=syllabus.consultation_room,
            consultation_contact=syllabus.consultation_contact,
            
            course_description=syllabus.course_description, 
            course_requirements=syllabus.course_requirements,
            
            version=1,
            status="Draft",
            
            chair_submitted_at=None,
            chair_rejected_at=None, 
            dean_submitted_at=None,
            dean_rejected_at=None,
            dean_approved_at=None,

            dean=dean_json,
            chair=chair_json,
        )

        # ✅ Clone ManyToMany fields (PEOs and Program Outcomes)
        new_syllabus.peos.set(PEO.objects.filter(program_id=syllabus.program.id, is_active=True))
        new_syllabus.program_outcomes.set(ProgramOutcome.objects.filter(program_id=syllabus.program.id, is_active=True)) 

        # Get only ACTIVE Program Outcomes for the program
        active_pos = ProgramOutcome.objects.filter(
            program=syllabus.program,
            is_active=True
        )

        active_po_ids = set(active_pos.values_list("id", flat=True)) 
        
        # ✅ Clone course outcomes
        old_outcomes = SyllabusCourseOutcome.objects.filter(syllabus=syllabus)
        for old_outcome in old_outcomes:
            new_outcome = SyllabusCourseOutcome.objects.create(
                syllabus=new_syllabus,
                co_code=old_outcome.co_code,
                co_description=old_outcome.co_description,
            )  
            
            # Get old COPO mappings
            old_copos = SyllCoPo.objects.filter(course_outcome=old_outcome)  
            for old_mapping in old_copos:
                old_po_id = old_mapping.program_outcome.id 
                # Only copy mapping if PO still exists and is active
                if old_po_id in active_po_ids:
                    SyllCoPo.objects.create(
                        syllabus=new_syllabus,
                        course_outcome=new_outcome,
                        program_outcome_id=old_po_id,
                        syllabus_co_po_code=old_mapping.syllabus_co_po_code,
                    ) 

        # ✅ Clone course outlines
        old_outlines = SyllabusCourseOutline.objects.filter(syllabus=syllabus)
        for old_outline in old_outlines:
            
            new_outline = SyllabusCourseOutline.objects.create(
                syllabus=new_syllabus,
                syllabus_term=old_outline.syllabus_term,
                row_no=old_outline.row_no, 
                allotted_hour=old_outline.allotted_hour,
                allotted_time=old_outline.allotted_time,
                intended_learning=old_outline.intended_learning,
                topics=old_outline.topics,
                suggested_readings=old_outline.suggested_readings,
                learning_activities=old_outline.learning_activities,
                assessment_tools=old_outline.assessment_tools,
                grading_criteria=old_outline.grading_criteria,
                remarks=old_outline.remarks,
            )

            # === SMART COT−CO MAPPING ===
            old_cotcos = SyllabusCotCo.objects.filter(course_outline=old_outline)
            for old_cotco in old_cotcos:
                old_co_code = old_cotco.course_outcome.co_code  # use co_code instead of ID

                # Find the new CO by matching co_code in new syllabus
                new_co = SyllabusCourseOutcome.objects.filter(
                    syllabus=new_syllabus,
                    co_code=old_co_code
                ).first()

                if not new_co:
                    # Skip or log instead of crashing
                    print(f"Warning: No matching CO for co_code={old_co_code}, skipping.")
                    continue

                SyllabusCotCo.objects.create( 
                    course_outline=new_outline,
                    course_outcome=new_co,
                ) 

        # Create another Report record for Dean Reports (version 1)
        Report.objects.create(
            syllabus=new_syllabus,
            bayanihan_group=new_syllabus.bayanihan_group, 
            version=new_syllabus.version,
        )
                
        new_syllabus_data = SyllabusDetailSerializer(new_syllabus, context={"request": request}).data
        return Response(new_syllabus_data, status=status.HTTP_201_CREATED)
    
    @action(detail=False, methods=["get"], url_path="approved-syllabi")
    def approved_syllabi(self, request): 
        user = request.user 
        role = request.query_params.get("role")     

        if not role:
            raise PermissionDenied("Role parameter is required.")
        role = role.upper()  # normalize
        
        # --- ADMIN ---
        if role == "ADMIN":
            if not UserRole.objects.filter(user=user, role__name="ADMIN").exists():
                raise PermissionDenied("You are not an Admin.")

            qs = (
                Syllabus.objects
                .filter(status="Approved by Dean", dean_approved_at__isnull=False)
                .select_related("course", "bayanihan_group")
                .prefetch_related("course_outlines")
                .order_by("bayanihan_group__school_year")
            )

        # --- BAYANIHAN LEADER ---
        elif role == "BAYANIHAN_LEADER":
            if not UserRole.objects.filter(user=user, role__name="BAYANIHAN_LEADER").exists():
                raise PermissionDenied("You are not a Bayanihan Leader.")

            leader_group_ids = BayanihanGroupUser.objects.filter(
                user=user, role="LEADER"
            ).values_list("group_id", flat=True)

            if not leader_group_ids:
                return Response(
                    {"detail": "You are not a leader in any Bayanihan group."},
                    status=403
                )

            qs = (
                Syllabus.objects
                .filter(
                    bayanihan_group_id__in=leader_group_ids,
                    status="Approved by Dean",
                    dean_approved_at__isnull=False
                )
                .select_related("course", "bayanihan_group")
                .prefetch_related("course_outlines")
                .order_by("bayanihan_group__school_year")
            )

        else:
            raise PermissionDenied("Invalid role parameter.")
        
        # --- FILTER OUT syllabi with both MIDTERM and FINALS TOS ---
        qs = qs.annotate(
            midterm_count=Count("tos_records", filter=Q(tos_records__term="MIDTERM")),
            finals_count=Count("tos_records", filter=Q(tos_records__term="FINALS")),
        ).filter(
            Q(midterm_count=0) | Q(finals_count=0)  # keep only if at least one is missing
        ).order_by("bayanihan_group__school_year")
 
        serializer = SyllabusApprovedReadSerializer(qs, many=True) 
        return Response(serializer.data, status=status.HTTP_200_OK)
    
    @action(detail=True, methods=["patch"], url_path="update-dates")
    def update_dates(self, request, pk=None):
        syllabus = self.get_object()
        serializer = SyllabusVersionSerializer(
            syllabus, data=request.data, partial=True
        )
        serializer.is_valid(raise_exception=True)
        serializer.save()
        return Response(serializer.data, status=status.HTTP_200_OK)
    
    @action(detail=True, methods=["get"])
    def export_docx(self, request, pk=None): 
        is_windows = platform.system().lower() == "windows" 
        if is_windows and pythoncom is not None:
            pythoncom.CoInitialize() 
           
        # Converts signature to absolute urls (https:127.0.0.1:8000/media/signatures/signature_name.png)
        def url_to_local_file(url):
            """Convert absolute URL to local MEDIA_ROOT path."""
            if not url:
                return None

            response = requests.get(url)
            if response.status_code != 200:
                return None

            # Create a temporary file
            suffix = os.path.splitext(url)[1] or ".png"
            tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
            tmp_file.write(response.content)
            tmp_file.flush()
            tmp_file.close()
            return tmp_file.name
         
        # Creates Course Outcomes Table in the Syllabus
        def build_copo_table(doc, syllabus):
            subdoc = Subdoc(doc)

            po_list = list(syllabus.program_outcomes.all())
            co_list = list(syllabus.course_outcomes.all())

            # Create header row
            table = subdoc.add_table(rows=2, cols=1 + len(po_list))
            table.style = 'Table Grid'

            # -----------------------------
            # PAGE WIDTH
            # -----------------------------
            MAX_TABLE_WIDTH_INCH = 4    # usable page width
            MAX_TABLE_WIDTH = Inches(MAX_TABLE_WIDTH_INCH)

            # Desired widths (in inches)
            DESIRED_CO_INCH = 1.5 
            DESIRED_CODE_INCH = 0.1

            # Total desired width in inches
            desired_total_width_inch = DESIRED_CO_INCH + len(po_list) * DESIRED_CODE_INCH

            # -----------------------------
            # SCALE IF NEEDED
            # -----------------------------
            if desired_total_width_inch > MAX_TABLE_WIDTH_INCH:
                scale_factor = MAX_TABLE_WIDTH_INCH / desired_total_width_inch
            else:
                scale_factor = 1.0

            # Final widths (convert to EMUs & int)
            CO_WIDTH = int(Inches(DESIRED_CO_INCH * scale_factor)) 
            CODE_WIDTH = int(Inches(DESIRED_CODE_INCH * scale_factor))

            # -----------------------------
            # APPLY COLUMN WIDTHS
            # -----------------------------
            # 1st column: CO (big)
            table.columns[0].width = CO_WIDTH

            # Following columns:
            # Row 1 = PO letter (medium)
            # Row 2+ = CO-PO codes (tiny)
            for idx in range(1, len(table.columns)):
                table.columns[idx].width = CODE_WIDTH   # use code width for all body rows

            # Helper for cell formatting
            def set_cell_text(cell, text, bold=False, size=9, align_center=True, valign_center=True):
                cell.text = text
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if align_center else WD_ALIGN_PARAGRAPH.LEFT
                    run = paragraph.runs[0]
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(size)
                    run.bold = bold
                    rFonts = run._element.rPr.rFonts
                    rFonts.set(qn('w:eastAsia'), 'Times New Roman')

                # Vertical alignment
                tc_pr = cell._tc.get_or_add_tcPr()
                from docx.oxml import OxmlElement
                vAlign = OxmlElement('w:vAlign')
                vAlign.set(qn('w:val'), 'center')
                tc_pr.append(vAlign)

            # ----- HEADER ROWS -----
            # CO header
            co_cell = table.cell(0, 0)
            co_cell.merge(table.cell(1, 0))
            set_cell_text(co_cell, "Course Outcomes (CO)", bold=True)

            # PO header merged
            if po_list:
                po_cell = table.cell(0, 1)
                for idx in range(1, len(po_list)):
                    po_cell.merge(table.cell(0, idx + 1))
                set_cell_text(po_cell, "Program Outcomes (PO)", bold=True)

            # PO code letters
            for idx, po in enumerate(po_list):
                cell = table.cell(1, idx + 1)
                set_cell_text(cell, po.po_letter)

            # ----- BODY ROWS -----
            for co in co_list:
                row_cells = table.add_row().cells

                # CO description (left aligned)
                set_cell_text(
                    row_cells[0],
                    f"{co.co_code}: {co.co_description}",
                    align_center=False
                )

                # CO-PO codes
                for idx, po in enumerate(po_list):
                    codes = syllabus.syllcopos.filter(
                        program_outcome=po,
                        course_outcome=co
                    ).values_list("syllabus_co_po_code", flat=True)

                    cell_text = ", ".join(codes) if codes else ""
                    set_cell_text(row_cells[idx + 1], cell_text)

            return subdoc

        # Preprocesses Course Requirement HTML data
        def preprocess_html(raw_html: str) -> str:
            """Clean up quirks in HTML exported from Word/TinyMCE."""
            if not raw_html:
                return ""

            # Decode entities using html module
            cleaned = ihtml.unescape(raw_html)

            # Apply PHP-like replacements
            replacements = [
                ("andbull;", "•"),   # fix bullet entity
                ("&bull;", "•"),     # also normal bull
                ("&nbsp;", ""),      # strip non-breaking spaces
                ("&", "and"),        # replace stray ampersands
                ("/n", "<br/>"),     # normalize line breaks
                ("andrsquo;", "'"),  # right single quote
                ("andndash;", "-"),  # dash
            ]
            for old, new in replacements:
                cleaned = cleaned.replace(old, new)

            return cleaned
        
        # Iterates Course Requirements HTML data and converts it to Docx elements 
        def html_to_subdoc(doc, raw_html):
            """
            Convert TinyMCE / Word HTML into a python-docx Subdoc,
            preserving bold, italic, underline, lists, and tables.
            """
            subdoc = Subdoc(doc)
            if not raw_html:
                return subdoc

            soup = BeautifulSoup(raw_html, "html.parser")
            
            def get_padding_left_in_inches(elem):
                style = elem.get("style", "")
                match = re.search(r"padding-left\s*:\s*(\d+)px", style)
                if match:
                    px = int(match.group(1))
                    # Approx conversion: 96 px = 1 inch
                    return px / 96
                return 0

            def apply_cell_style(cell, td_elem):
                # Background color
                style = td_elem.attrs.get("style", "")
                match_bg = re.search(r'background-color\s*:\s*#?([0-9a-fA-F]{6})', style)
                if match_bg:
                    color = match_bg.group(1)
                    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                    
                # Text color
                match_color = re.search(r'color\s*:\s*#?([0-9a-fA-F]{6})', style)
                if match_color:
                    hex_color = match_color.group(1)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor.from_string(hex_color)

                # Text alignment
                align = re.search(r'text-align\s*:\s*(\w+)', style)
                if align:
                    if align.group(1).lower() == 'center':
                        for p in cell.paragraphs:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif align.group(1).lower() == 'right':
                        for p in cell.paragraphs:
                            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        for p in cell.paragraphs:
                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Vertical alignment
                valign = td_elem.attrs.get("valign", "").lower()
                if valign == "middle":
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                elif valign == "bottom":
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
                else:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

                # ----------------- Table cell border -----------------
                border_color = re.search(r'border-color\s*:\s*#?([0-9a-fA-F]{6})', style)
                border_width = re.search(r'border\s*:\s*(\d+)px', style)  # simple px parsing
                if border_color or border_width:
                    tc_pr = cell._tc.get_or_add_tcPr()
                    sz = int(border_width.group(1))*4 if border_width else 4
                    color = border_color.group(1) if border_color else "000000"

                    borders_xml = f'''
                    <w:tcBorders {nsdecls('w')}>
                    <w:top w:val="single" w:sz="{sz}" w:color="{color}"/>
                    <w:left w:val="single" w:sz="{sz}" w:color="{color}"/>
                    <w:bottom w:val="single" w:sz="{sz}" w:color="{color}"/>
                    <w:right w:val="single" w:sz="{sz}" w:color="{color}"/>
                    </w:tcBorders>
                    '''
                    tc_pr.append(parse_xml(borders_xml))
                    
            # ----------------- Helper: Recursive inline rendering -----------------
            def render_inline(parent_para, elem, bold=False, italic=False, underline=False, color=None):
                if isinstance(elem, NavigableString):
                    run = parent_para.add_run(str(elem))
                    run.bold = bold
                    run.italic = italic
                    run.underline = underline
                    run.font.size = Pt(10)
                    run.font.name = "Times New Roman"
                    rFonts = run._element.rPr.rFonts
                    rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                    if color:
                        run.font.color.rgb = RGBColor.from_string(color)
                elif hasattr(elem, "name"):
                    # Update formatting based on current tag
                    new_bold = bold or elem.name in ["b", "strong"]
                    new_italic = italic or elem.name in ["i", "em"]
                    new_underline = underline or elem.name == "u"

                    # Check inline style
                    style = elem.attrs.get("style", "").lower()
                    inline_color = None
                    match_color = re.search(r'color\s*:\s*#?([0-9a-fA-F]{6})', style)
                    if match_color:
                        inline_color = match_color.group(1)
                        
                    if "font-weight: bold" in style:
                        new_bold = True
                    if "font-style: italic" in style:
                        new_italic = True
                    if "text-decoration: underline" in style:
                        new_underline = True

                    if elem.name == "br":
                        parent_para.add_run().add_break()

                    for child in elem.children:
                        render_inline(parent_para, child, new_bold, new_italic, new_underline, inline_color or color)

            # ----------------- Helper: Process lists recursively -----------------
            def process_list(list_elem, level=0):
                items = list_elem.find_all("li", recursive=False)
                for idx, li in enumerate(items, start=1):
                    p = subdoc.add_paragraph()
                    # Indentation for nesting
                    p.paragraph_format.left_indent = Inches(get_padding_left_in_inches(li))
                    p.paragraph_format.first_line_indent = Inches(-0.2)
                    p.paragraph_format.space_after = Pt(0)

                    # Bullet or number prefix
                    prefix = "• " if list_elem.name == "ul" else f"{idx}. "
                    run = p.add_run(prefix)
                    run.font.size = Pt(10)

                    # Render li text + inline formatting recursively
                    for child in li.contents:
                        render_inline(p, child)

                    # Handle nested lists inside this li
                    for child_list in li.find_all(["ul", "ol"], recursive=False):
                        process_list(child_list, level + 1)

            # ----------------- Helper: Render table -----------------
            def render_table(table_elem):
                rows = table_elem.find_all("tr", recursive=True)
                if not rows:
                    return None

                # Calculate max columns considering colspan
                n_cols = 0
                for r in rows:
                    col_count = 0
                    for td in r.find_all(["td", "th"], recursive=False):
                        col_count += int(td.get("colspan", 1))
                    n_cols = max(n_cols, col_count)

                table = subdoc.add_table(rows=len(rows), cols=n_cols)
                table.style = "Table Grid"
                # Remove table border if specified
                table_style = table_elem.get("style", "").lower()
                if "border:none" in table_style or "border-width:0" in table_style:
                    tbl = table._tbl
                    tblPr = tbl.tblPr
                    tblBorders = tblPr.tblBorders
                    if tblBorders is not None:
                        tblPr.remove(tblBorders)

                # Track merged cells positions
                merge_map = [[None for _ in range(n_cols)] for _ in range(len(rows))]

                for r_idx, row in enumerate(rows):
                    c_idx = 0
                    for td in row.find_all(["td", "th"], recursive=False):
                        # Find next available column (skip merged cells)
                        while c_idx < n_cols and merge_map[r_idx][c_idx]:
                            c_idx += 1

                        colspan = int(td.get("colspan", 1))
                        rowspan = int(td.get("rowspan", 1))
                        start_cell = table.cell(r_idx, c_idx)
                        start_cell.text = ""

                        apply_cell_style(start_cell, td)
                        para = start_cell.paragraphs[0]

                        # Render inline content inside cell
                        for child in td.children:
                            if getattr(child, "name", None) == "table":
                                render_table(child)  # nested table
                            else:
                                render_inline(para, child)

                        # Merge cells if colspan or rowspan > 1
                        if colspan > 1 or rowspan > 1:
                            end_row = r_idx + rowspan - 1
                            end_col = c_idx + colspan - 1
                            end_cell = table.cell(end_row, end_col)
                            start_cell.merge(end_cell)

                            # Mark merged cells in the map
                            for mr in range(r_idx, r_idx + rowspan):
                                for mc in range(c_idx, c_idx + colspan):
                                    merge_map[mr][mc] = True # type: ignore

                        c_idx += colspan
                return table 

            # ----------------- Main loop over top-level elements -----------------
            for elem in soup.children:
                if isinstance(elem, NavigableString):
                    if str(elem).strip():
                        p = subdoc.add_paragraph(str(elem).strip())
                elif getattr(elem, "name", None) == "p":
                    p = subdoc.add_paragraph()
                    indent = Inches(get_padding_left_in_inches(elem))
                    p.paragraph_format.left_indent = indent
                    p.paragraph_format.first_line_indent = Inches(-0.2 if str(elem).strip().startswith("●") else 0)
                    p.paragraph_format.space_after = Pt(0)

                    # Add bullet prefix if the paragraph starts with "●"
                    text_content = "".join(str(c) for c in getattr(elem, "contents", []))
                    if text_content.strip().startswith("●"):
                        run = p.add_run("• ")
                        run.font.size = Pt(10)
                        # remove bullet character from text
                        elem.contents[0].replace_with(str(elem.contents[0]).replace("●", "", 1)) # type: ignore

                    for child in getattr(elem, "contents", []):
                        render_inline(p, child)
                elif getattr(elem, "name", None) in ["ul", "ol"]:
                    process_list(elem)
                elif getattr(elem, "name", None) == "table":
                    render_table(elem)

            return subdoc

        try:
            syllabus = self.get_object() 

            # 1. Load your DOCX template
            template_path = os.path.join(settings.BASE_DIR, "syllabi", "templates", "SyllabusTemp.docx")
            doc = DocxTemplate(template_path)
            
            # ---------- Build CO–PO Table as HTML ----------  
            copo_subdoc = build_copo_table(doc, syllabus)

            # ---------- Build Course Outline Data ----------
            outline_rows_mid = []
            outline_rows_final = []     
            
            for outline in syllabus.course_outlines.all():
                # Collect all CO codes linked to this outline
                co_codes = ", ".join(
                    cotco.course_outcome.co_code
                    for cotco in outline.cotcos.select_related("course_outcome").all()
                )
                row = {
                    "allotted_time": f"{outline.allotted_hour or ''} hours, {outline.allotted_time or ''}".strip(),
                    "co_code": co_codes or "",
                    "ilo": outline.intended_learning or "",
                    "topics": outline.topics or "",
                    "readings": outline.suggested_readings or "",
                    "activities": outline.learning_activities or "",
                    "assessment": outline.assessment_tools or "",
                    "grading": outline.grading_criteria or "",
                    "remarks": outline.remarks or "",
                }

                if outline.syllabus_term == "MIDTERM":
                    outline_rows_mid.append(row)
                elif outline.syllabus_term == "FINALS":
                    outline_rows_final.append(row)   
    
            # ---------- Build Course Requirements Data ----------
            req_html = syllabus.course_requirements or ""
            course_requirements_subdoc = html_to_subdoc(doc, req_html)  
            
            # --- Bayanihan Leaders ----   
            leaders = [
                m for m in syllabus.bayanihan_group.bayanihan_members.all()
                if m.role == "LEADER"
            ]
            leaders_context = []
            for m in leaders:
                sig_url = getattr(m.user, "signature", None) 
                clean_url = get_public_url(sig_url.url) if sig_url else None
                
                local_sig_path = url_to_local_file(clean_url) if clean_url else None
                    
                # sig_path = sig_url.url if sig_url else None
                leaders_context.append({
                    "fname": m.user.first_name,
                    "lname": m.user.last_name,
                    "signature": InlineImage(doc, local_sig_path , width=Mm(20)) if local_sig_path  else "",
                })   

            # --- Chairperson ---
            chair_name = syllabus.chair.get("name") if syllabus.chair else ""
            chair_sig_url = syllabus.chair.get("signature") if syllabus.chair else None
            chair_sig_path = url_to_local_file(chair_sig_url) if chair_sig_url else None
            chair_signature = InlineImage(doc, chair_sig_path, width=Mm(20)) if chair_sig_path else ""

            # --- Dean ---
            dean_name = syllabus.dean.get("name") if syllabus.dean else ""
            dean_sig_url = syllabus.dean.get("signature") if syllabus.dean else None
            dean_sig_path = url_to_local_file(dean_sig_url) if dean_sig_url else None
            dean_signature = InlineImage(doc, dean_sig_path, width=Mm(20)) if dean_sig_path else ""

            # 2. Prepare context with syllabus fields
            context = {
                "version": (
                    f"{int(syllabus.syllabus_template.revision_no):02d}"
                    if syllabus.syllabus_template and syllabus.syllabus_template.revision_no is not None
                    else ""
                ),
                "effective_date": syllabus.effective_date.strftime("%m.%d.%y") if syllabus.effective_date else "",
                "college_description": syllabus.college.college_description if syllabus.college else "",
                "department_name": syllabus.program.department.department_name if syllabus.program.department else "",
                "peos": syllabus.peos.all(),
                "pos": syllabus.program_outcomes.all(),
                "course_title": syllabus.course.course_title,
                "course_code": syllabus.course.course_code,
                "course_semester": syllabus.course.course_semester.lower(),
                "school_year": syllabus.bayanihan_group.school_year,  
                "course_description": syllabus.course_description,
                "class_schedules": syllabus.class_schedules,
                "building_room": syllabus.building_room,
                "course_pre_req": syllabus.course.course_pre_req,
                "course_co_req": syllabus.course.course_co_req,
                "consultation_hours": syllabus.consultation_hours,
                "consultation_room": syllabus.consultation_room,
                "consultation_contact": syllabus.consultation_contact,
                "instructor_names": ", ".join(i.user.get_full_name() for i in syllabus.instructors.all()),
                "instructor_emails": ", ".join(i.user.email for i in syllabus.instructors.all()), 
                "class_contact": syllabus.class_contact,
                "course_credit_unit": syllabus.course.course_credit_unit,
                "course_unit_lec": syllabus.course.course_unit_lec,
                "course_unit_lab": syllabus.course.course_unit_lab,
                "course_semester": syllabus.course.course_semester,

                # Special placeholders
                "copo": copo_subdoc,
                "course_outlines_mid": outline_rows_mid, 
                "course_outlines_final": outline_rows_final, 
                "course_requirements": course_requirements_subdoc,
                
                # Signatories
                "leaders": leaders_context, 
                "chair": chair_name,
                "chair_signature": chair_signature,
                "dept": syllabus.program.department.department_code if syllabus.program.department else "",
                "dean": dean_name,
                "dean_signature": dean_signature,
                "college": syllabus.college.college_code if syllabus.college else "",
            }

            # 3. Render the DOCX with data
            doc.render(context)
            
            timestamp = syllabus.dean_approved_at.strftime("%Y-%m-%d") if syllabus.dean_approved_at else "NA"
            filename_docx = f"Syllabus-{syllabus.course.course_code}-{syllabus.course.course_semester.lower()} Semester-{syllabus.bayanihan_group.school_year}-{syllabus.status}-{timestamp}.docx"
            
            file_bytes = io.BytesIO()
            doc.save(file_bytes)
            file_bytes.seek(0)

            docx_path = default_storage.save(
                f"syllabi/{filename_docx}",
                ContentFile(file_bytes.read())
            )

            docx_url = default_storage.url(docx_path)

            return JsonResponse({"docx_url": docx_url})
        
            # filename_pdf = filename_docx.replace(".docx", ".pdf")  

            # # Use a temp directory so we can run LibreOffice/win32com conversion
            # with tempfile.TemporaryDirectory() as tmpdir:
            #     temp_docx = os.path.join(tmpdir, filename_docx)
            #     temp_pdf = os.path.join(tmpdir, filename_pdf)

            #     # Save generated .docx locally for conversion
            #     doc.save(temp_docx)

            #     # Convert DOCX -> PDF
            #     convert_docx_to_pdf(temp_docx, temp_pdf)

            #     # Read PDF file
            #     with open(temp_pdf, "rb") as f:
            #         pdf_bytes = f.read()

            #     # Upload to DigitalOcean Spaces
            #     pdf_path = default_storage.save(
            #         f"syllabi/{filename_pdf}",  # folder inside bucket
            #         ContentFile(pdf_bytes)
            #     )

            # # Get the public URL for the browser to download
            # pdf_url = default_storage.url(pdf_path)
 
            # # 9️⃣ Return PDF as HTTP response
            # return JsonResponse({"pdf_url": pdf_url}) 
            # return FileResponse(open(output_pdf, "rb"), content_type="application/pdf") 

            # # 4. Convert DOCX → PDF
            # convert_docx_to_pdf(output_docx, output_pdf)  # works on Windows
            # # Alternative (Linux/macOS): os.system(f"libreoffice --headless --convert-to pdf {output_docx} --outdir {settings.MEDIA_ROOT}")
            
            # # 5. Return the PDF file as HTTP response
            # return FileResponse(open(output_pdf, "rb"), content_type="application/pdf") 
            
        finally:
            if is_windows and pythoncom is not None:
                pythoncom.CoUninitialize()

    @action(detail=True, methods=["get"])
    def export_pdf(self, request, pk=None): 
        is_windows = platform.system().lower() == "windows" 
        if is_windows and pythoncom is not None:
            pythoncom.CoInitialize() 
           
        # Converts signature to absolute urls (https:127.0.0.1:8000/media/signatures/signature_name.png)
        def url_to_local_file(url):
            """Convert absolute URL to local MEDIA_ROOT path."""
            if not url:
                return None

            response = requests.get(url)
            if response.status_code != 200:
                return None

            # Create a temporary file
            suffix = os.path.splitext(url)[1] or ".png"
            tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
            tmp_file.write(response.content)
            tmp_file.flush()
            tmp_file.close()
            return tmp_file.name
         
        # Creates Course Outcomes Table in the Syllabus
        def build_copo_table(doc, syllabus):
            subdoc = Subdoc(doc)

            po_list = list(syllabus.program_outcomes.all())
            co_list = list(syllabus.course_outcomes.all())

            # Create header row
            table = subdoc.add_table(rows=2, cols=1 + len(po_list))
            table.style = 'Table Grid'

            # -----------------------------
            # PAGE WIDTH
            # -----------------------------
            MAX_TABLE_WIDTH_INCH = 4    # usable page width
            MAX_TABLE_WIDTH = Inches(MAX_TABLE_WIDTH_INCH)

            # Desired widths (in inches)
            DESIRED_CO_INCH = 1.5 
            DESIRED_CODE_INCH = 0.1

            # Total desired width in inches
            desired_total_width_inch = DESIRED_CO_INCH + len(po_list) * DESIRED_CODE_INCH

            # -----------------------------
            # SCALE IF NEEDED
            # -----------------------------
            if desired_total_width_inch > MAX_TABLE_WIDTH_INCH:
                scale_factor = MAX_TABLE_WIDTH_INCH / desired_total_width_inch
            else:
                scale_factor = 1.0

            # Final widths (convert to EMUs & int)
            CO_WIDTH = int(Inches(DESIRED_CO_INCH * scale_factor)) 
            CODE_WIDTH = int(Inches(DESIRED_CODE_INCH * scale_factor))

            # -----------------------------
            # APPLY COLUMN WIDTHS
            # -----------------------------
            # 1st column: CO (big)
            table.columns[0].width = CO_WIDTH

            # Following columns:
            # Row 1 = PO letter (medium)
            # Row 2+ = CO-PO codes (tiny)
            for idx in range(1, len(table.columns)):
                table.columns[idx].width = CODE_WIDTH   # use code width for all body rows

            # Helper for cell formatting
            def set_cell_text(cell, text, bold=False, size=9, align_center=True, valign_center=True):
                cell.text = text
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if align_center else WD_ALIGN_PARAGRAPH.LEFT
                    run = paragraph.runs[0]
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(size)
                    run.bold = bold
                    rFonts = run._element.rPr.rFonts
                    rFonts.set(qn('w:eastAsia'), 'Times New Roman')

                # Vertical alignment
                tc_pr = cell._tc.get_or_add_tcPr()
                from docx.oxml import OxmlElement
                vAlign = OxmlElement('w:vAlign')
                vAlign.set(qn('w:val'), 'center')
                tc_pr.append(vAlign)

            # ----- HEADER ROWS -----
            # CO header
            co_cell = table.cell(0, 0)
            co_cell.merge(table.cell(1, 0))
            set_cell_text(co_cell, "Course Outcomes (CO)", bold=True)

            # PO header merged
            if po_list:
                po_cell = table.cell(0, 1)
                for idx in range(1, len(po_list)):
                    po_cell.merge(table.cell(0, idx + 1))
                set_cell_text(po_cell, "Program Outcomes (PO)", bold=True)

            # PO code letters
            for idx, po in enumerate(po_list):
                cell = table.cell(1, idx + 1)
                set_cell_text(cell, po.po_letter)

            # ----- BODY ROWS -----
            for co in co_list:
                row_cells = table.add_row().cells

                # CO description (left aligned)
                set_cell_text(
                    row_cells[0],
                    f"{co.co_code}: {co.co_description}",
                    align_center=False
                )

                # CO-PO codes
                for idx, po in enumerate(po_list):
                    codes = syllabus.syllcopos.filter(
                        program_outcome=po,
                        course_outcome=co
                    ).values_list("syllabus_co_po_code", flat=True)

                    cell_text = ", ".join(codes) if codes else ""
                    set_cell_text(row_cells[idx + 1], cell_text)

            return subdoc

        # Preprocesses Course Requirement HTML data
        def preprocess_html(raw_html: str) -> str:
            """Clean up quirks in HTML exported from Word/TinyMCE."""
            if not raw_html:
                return ""

            # Decode entities using html module
            cleaned = ihtml.unescape(raw_html)

            # Apply PHP-like replacements
            replacements = [
                ("andbull;", "•"),   # fix bullet entity
                ("&bull;", "•"),     # also normal bull
                ("&nbsp;", ""),      # strip non-breaking spaces
                ("&", "and"),        # replace stray ampersands
                ("/n", "<br/>"),     # normalize line breaks
                ("andrsquo;", "'"),  # right single quote
                ("andndash;", "-"),  # dash
            ]
            for old, new in replacements:
                cleaned = cleaned.replace(old, new)

            return cleaned
        
        # Iterates Course Requirements HTML data and converts it to Docx elements 
        def html_to_subdoc(doc, raw_html):
            """
            Convert TinyMCE / Word HTML into a python-docx Subdoc,
            preserving bold, italic, underline, lists, and tables.
            """
            subdoc = Subdoc(doc)
            if not raw_html:
                return subdoc

            soup = BeautifulSoup(raw_html, "html.parser")
            
            def get_padding_left_in_inches(elem):
                style = elem.get("style", "")
                match = re.search(r"padding-left\s*:\s*(\d+)px", style)
                if match:
                    px = int(match.group(1))
                    # Approx conversion: 96 px = 1 inch
                    return px / 96
                return 0

            def apply_cell_style(cell, td_elem):
                # Background color
                style = td_elem.attrs.get("style", "")
                match_bg = re.search(r'background-color\s*:\s*#?([0-9a-fA-F]{6})', style)
                if match_bg:
                    color = match_bg.group(1)
                    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                    
                # Text color
                match_color = re.search(r'color\s*:\s*#?([0-9a-fA-F]{6})', style)
                if match_color:
                    hex_color = match_color.group(1)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor.from_string(hex_color)

                # Text alignment
                align = re.search(r'text-align\s*:\s*(\w+)', style)
                if align:
                    if align.group(1).lower() == 'center':
                        for p in cell.paragraphs:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif align.group(1).lower() == 'right':
                        for p in cell.paragraphs:
                            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        for p in cell.paragraphs:
                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Vertical alignment
                valign = td_elem.attrs.get("valign", "").lower()
                if valign == "middle":
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                elif valign == "bottom":
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
                else:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

                # ----------------- Table cell border -----------------
                border_color = re.search(r'border-color\s*:\s*#?([0-9a-fA-F]{6})', style)
                border_width = re.search(r'border\s*:\s*(\d+)px', style)  # simple px parsing
                if border_color or border_width:
                    tc_pr = cell._tc.get_or_add_tcPr()
                    sz = int(border_width.group(1))*4 if border_width else 4
                    color = border_color.group(1) if border_color else "000000"

                    borders_xml = f'''
                    <w:tcBorders {nsdecls('w')}>
                    <w:top w:val="single" w:sz="{sz}" w:color="{color}"/>
                    <w:left w:val="single" w:sz="{sz}" w:color="{color}"/>
                    <w:bottom w:val="single" w:sz="{sz}" w:color="{color}"/>
                    <w:right w:val="single" w:sz="{sz}" w:color="{color}"/>
                    </w:tcBorders>
                    '''
                    tc_pr.append(parse_xml(borders_xml))
                    
            # ----------------- Helper: Recursive inline rendering -----------------
            def render_inline(parent_para, elem, bold=False, italic=False, underline=False, color=None):
                if isinstance(elem, NavigableString):
                    run = parent_para.add_run(str(elem))
                    run.bold = bold
                    run.italic = italic
                    run.underline = underline
                    run.font.size = Pt(10)
                    run.font.name = "Times New Roman"
                    rFonts = run._element.rPr.rFonts
                    rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                    if color:
                        run.font.color.rgb = RGBColor.from_string(color)
                elif hasattr(elem, "name"):
                    # Update formatting based on current tag
                    new_bold = bold or elem.name in ["b", "strong"]
                    new_italic = italic or elem.name in ["i", "em"]
                    new_underline = underline or elem.name == "u"

                    # Check inline style
                    style = elem.attrs.get("style", "").lower()
                    inline_color = None
                    match_color = re.search(r'color\s*:\s*#?([0-9a-fA-F]{6})', style)
                    if match_color:
                        inline_color = match_color.group(1)
                        
                    if "font-weight: bold" in style:
                        new_bold = True
                    if "font-style: italic" in style:
                        new_italic = True
                    if "text-decoration: underline" in style:
                        new_underline = True

                    if elem.name == "br":
                        parent_para.add_run().add_break()

                    for child in elem.children:
                        render_inline(parent_para, child, new_bold, new_italic, new_underline, inline_color or color)

            # ----------------- Helper: Process lists recursively -----------------
            def process_list(list_elem, level=0):
                items = list_elem.find_all("li", recursive=False)
                for idx, li in enumerate(items, start=1):
                    p = subdoc.add_paragraph()
                    # Indentation for nesting
                    p.paragraph_format.left_indent = Inches(get_padding_left_in_inches(li))
                    p.paragraph_format.first_line_indent = Inches(-0.2)
                    p.paragraph_format.space_after = Pt(0)

                    # Bullet or number prefix
                    prefix = "• " if list_elem.name == "ul" else f"{idx}. "
                    run = p.add_run(prefix)
                    run.font.size = Pt(10)

                    # Render li text + inline formatting recursively
                    for child in li.contents:
                        render_inline(p, child)

                    # Handle nested lists inside this li
                    for child_list in li.find_all(["ul", "ol"], recursive=False):
                        process_list(child_list, level + 1)

            # ----------------- Helper: Render table -----------------
            def render_table(table_elem):
                rows = table_elem.find_all("tr", recursive=True)
                if not rows:
                    return None

                # Calculate max columns considering colspan
                n_cols = 0
                for r in rows:
                    col_count = 0
                    for td in r.find_all(["td", "th"], recursive=False):
                        col_count += int(td.get("colspan", 1))
                    n_cols = max(n_cols, col_count)

                table = subdoc.add_table(rows=len(rows), cols=n_cols)
                table.style = "Table Grid"
                # Remove table border if specified
                table_style = table_elem.get("style", "").lower()
                if "border:none" in table_style or "border-width:0" in table_style:
                    tbl = table._tbl
                    tblPr = tbl.tblPr
                    tblBorders = tblPr.tblBorders
                    if tblBorders is not None:
                        tblPr.remove(tblBorders)

                # Track merged cells positions
                merge_map = [[None for _ in range(n_cols)] for _ in range(len(rows))]

                for r_idx, row in enumerate(rows):
                    c_idx = 0
                    for td in row.find_all(["td", "th"], recursive=False):
                        # Find next available column (skip merged cells)
                        while c_idx < n_cols and merge_map[r_idx][c_idx]:
                            c_idx += 1

                        colspan = int(td.get("colspan", 1))
                        rowspan = int(td.get("rowspan", 1))
                        start_cell = table.cell(r_idx, c_idx)
                        start_cell.text = ""

                        apply_cell_style(start_cell, td)
                        para = start_cell.paragraphs[0]

                        # Render inline content inside cell
                        for child in td.children:
                            if getattr(child, "name", None) == "table":
                                render_table(child)  # nested table
                            else:
                                render_inline(para, child)

                        # Merge cells if colspan or rowspan > 1
                        if colspan > 1 or rowspan > 1:
                            end_row = r_idx + rowspan - 1
                            end_col = c_idx + colspan - 1
                            end_cell = table.cell(end_row, end_col)
                            start_cell.merge(end_cell)

                            # Mark merged cells in the map
                            for mr in range(r_idx, r_idx + rowspan):
                                for mc in range(c_idx, c_idx + colspan):
                                    merge_map[mr][mc] = True # type: ignore

                        c_idx += colspan
                return table 

            # ----------------- Main loop over top-level elements -----------------
            for elem in soup.children:
                if isinstance(elem, NavigableString):
                    if str(elem).strip():
                        p = subdoc.add_paragraph(str(elem).strip())
                elif getattr(elem, "name", None) == "p":
                    p = subdoc.add_paragraph()
                    indent = Inches(get_padding_left_in_inches(elem))
                    p.paragraph_format.left_indent = indent
                    p.paragraph_format.first_line_indent = Inches(-0.2 if str(elem).strip().startswith("●") else 0)
                    p.paragraph_format.space_after = Pt(0)

                    # Add bullet prefix if the paragraph starts with "●"
                    text_content = "".join(str(c) for c in getattr(elem, "contents", []))
                    if text_content.strip().startswith("●"):
                        run = p.add_run("• ")
                        run.font.size = Pt(10)
                        # remove bullet character from text
                        elem.contents[0].replace_with(str(elem.contents[0]).replace("●", "", 1)) # type: ignore

                    for child in getattr(elem, "contents", []):
                        render_inline(p, child)
                elif getattr(elem, "name", None) in ["ul", "ol"]:
                    process_list(elem)
                elif getattr(elem, "name", None) == "table":
                    render_table(elem)

            return subdoc

        try:
            syllabus = self.get_object() 

            # 1. Load your DOCX template
            template_path = os.path.join(settings.BASE_DIR, "syllabi", "templates", "SyllabusTemp.docx")
            doc = DocxTemplate(template_path)
            
            # ---------- Build CO–PO Table as HTML ----------  
            copo_subdoc = build_copo_table(doc, syllabus)

            # ---------- Build Course Outline Data ----------
            outline_rows_mid = []
            outline_rows_final = []     
            
            for outline in syllabus.course_outlines.all():
                # Collect all CO codes linked to this outline
                co_codes = ", ".join(
                    cotco.course_outcome.co_code
                    for cotco in outline.cotcos.select_related("course_outcome").all()
                )
                row = {
                    "allotted_time": f"{outline.allotted_hour or ''} hours, {outline.allotted_time or ''}".strip(),
                    "co_code": co_codes or "",
                    "ilo": outline.intended_learning or "",
                    "topics": outline.topics or "",
                    "readings": outline.suggested_readings or "",
                    "activities": outline.learning_activities or "",
                    "assessment": outline.assessment_tools or "",
                    "grading": outline.grading_criteria or "",
                    "remarks": outline.remarks or "",
                }

                if outline.syllabus_term == "MIDTERM":
                    outline_rows_mid.append(row)
                elif outline.syllabus_term == "FINALS":
                    outline_rows_final.append(row)   
    
            # ---------- Build Course Requirements Data ----------
            req_html = syllabus.course_requirements or ""
            course_requirements_subdoc = html_to_subdoc(doc, req_html)  
            
            # --- Bayanihan Leaders ----   
            leaders = [
                m for m in syllabus.bayanihan_group.bayanihan_members.all()
                if m.role == "LEADER"
            ]
            leaders_context = []
            for m in leaders:
                sig_url = getattr(m.user, "signature", None) 
                clean_url = get_public_url(sig_url.url) if sig_url else None
                
                local_sig_path = url_to_local_file(clean_url) if clean_url else None
                    
                # sig_path = sig_url.url if sig_url else None
                leaders_context.append({
                    "fname": m.user.first_name,
                    "lname": m.user.last_name,
                    "signature": InlineImage(doc, local_sig_path , width=Mm(20)) if local_sig_path  else "",
                })   

            # --- Chairperson ---
            chair_name = syllabus.chair.get("name") if syllabus.chair else ""
            chair_sig_url = syllabus.chair.get("signature") if syllabus.chair else None
            chair_sig_path = url_to_local_file(chair_sig_url) if chair_sig_url else None
            chair_signature = InlineImage(doc, chair_sig_path, width=Mm(20)) if chair_sig_path else ""

            # --- Dean ---
            dean_name = syllabus.dean.get("name") if syllabus.dean else ""
            dean_sig_url = syllabus.dean.get("signature") if syllabus.dean else None
            dean_sig_path = url_to_local_file(dean_sig_url) if dean_sig_url else None
            dean_signature = InlineImage(doc, dean_sig_path, width=Mm(20)) if dean_sig_path else ""

            # 2. Prepare context with syllabus fields
            context = {
                "version": (
                    f"{int(syllabus.syllabus_template.revision_no):02d}"
                    if syllabus.syllabus_template and syllabus.syllabus_template.revision_no is not None
                    else ""
                ),
                "effective_date": syllabus.effective_date.strftime("%m.%d.%y") if syllabus.effective_date else "",
                "college_description": syllabus.college.college_description if syllabus.college else "",
                "department_name": syllabus.program.department.department_name if syllabus.program.department else "",
                "peos": syllabus.peos.all(),
                "pos": syllabus.program_outcomes.all(),
                "course_title": syllabus.course.course_title,
                "course_code": syllabus.course.course_code,
                "course_semester": syllabus.course.course_semester.lower(),
                "school_year": syllabus.bayanihan_group.school_year,  
                "course_description": syllabus.course_description,
                "class_schedules": syllabus.class_schedules,
                "building_room": syllabus.building_room,
                "course_pre_req": syllabus.course.course_pre_req,
                "course_co_req": syllabus.course.course_co_req,
                "consultation_hours": syllabus.consultation_hours,
                "consultation_room": syllabus.consultation_room,
                "consultation_contact": syllabus.consultation_contact,
                "instructor_names": ", ".join(i.user.get_full_name() for i in syllabus.instructors.all()),
                "instructor_emails": ", ".join(i.user.email for i in syllabus.instructors.all()), 
                "class_contact": syllabus.class_contact,
                "course_credit_unit": syllabus.course.course_credit_unit,
                "course_unit_lec": syllabus.course.course_unit_lec,
                "course_unit_lab": syllabus.course.course_unit_lab,
                "course_semester": syllabus.course.course_semester,

                # Special placeholders
                "copo": copo_subdoc,
                "course_outlines_mid": outline_rows_mid, 
                "course_outlines_final": outline_rows_final, 
                "course_requirements": course_requirements_subdoc,
                
                # Signatories
                "leaders": leaders_context, 
                "chair": chair_name,
                "chair_signature": chair_signature,
                "dept": syllabus.program.department.department_code if syllabus.program.department else "",
                "dean": dean_name,
                "dean_signature": dean_signature,
                "college": syllabus.college.college_code if syllabus.college else "",
            }

            # 3. Render the DOCX with data
            doc.render(context)
            
            timestamp = syllabus.dean_approved_at.strftime("%Y-%m-%d") if syllabus.dean_approved_at else "NA"
            filename_docx = f"Syllabus-{syllabus.course.course_code}-{syllabus.course.course_semester.lower()} Semester-{syllabus.bayanihan_group.school_year}-{syllabus.status}-{timestamp}.docx"
        
            filename_pdf = filename_docx.replace(".docx", ".pdf")  

            # Use a temp directory so we can run LibreOffice/win32com conversion
            with tempfile.TemporaryDirectory() as tmpdir:
                temp_docx = os.path.join(tmpdir, filename_docx)
                temp_pdf = os.path.join(tmpdir, filename_pdf)

                # Save generated .docx locally for conversion
                doc.save(temp_docx)

                # Convert DOCX -> PDF
                convert_docx_to_pdf(temp_docx, temp_pdf)

                # Read PDF file
                with open(temp_pdf, "rb") as f:
                    pdf_bytes = f.read()

                # Upload to DigitalOcean Spaces
                pdf_path = default_storage.save(
                    f"syllabi/{filename_pdf}",  # folder inside bucket
                    ContentFile(pdf_bytes)
                )

            # Get the public URL for the browser to download
            pdf_url = default_storage.url(pdf_path)
 
            # 9️⃣ Return PDF as HTTP response
            return JsonResponse({"pdf_url": pdf_url}) 
            # return FileResponse(open(output_pdf, "rb"), content_type="application/pdf") 

            # # 4. Convert DOCX → PDF
            # convert_docx_to_pdf(output_docx, output_pdf)  # works on Windows
            # # Alternative (Linux/macOS): os.system(f"libreoffice --headless --convert-to pdf {output_docx} --outdir {settings.MEDIA_ROOT}")
            
            # # 5. Return the PDF file as HTTP response
            # return FileResponse(open(output_pdf, "rb"), content_type="application/pdf") 
            
        finally:
            if is_windows and pythoncom is not None:
                pythoncom.CoUninitialize()

class SyllabusCourseOutcomeViewSet(viewsets.ModelViewSet):
    serializer_class = SyllabusCourseOutcomeSerializer
    queryset = SyllabusCourseOutcome.objects.all()
    permission_classes = [RolePermission("ADMIN", "BAYANIHAN_LEADER")]

    def get_queryset(self): 
        syllabus_id = self.request.GET.get("syllabus_id")
        if syllabus_id:
            return SyllabusCourseOutcome.objects.filter(syllabus_id=syllabus_id) 
        return SyllabusCourseOutcome.objects.all()
    

class SyllCoPoViewSet(viewsets.ModelViewSet):
    serializer_class = SyllCoPoSerializer
    queryset = SyllCoPo.objects.all()
    permission_classes = [RolePermission("ADMIN", "BAYANIHAN_LEADER")]

    def get_queryset(self):
        syllabus_id = self.request.GET.get("syllabus_id")
        if syllabus_id:
            return SyllCoPo.objects.filter(syllabus_id=syllabus_id).select_related(
                "course_outcome", "program_outcome"
            )
        return SyllCoPo.objects.all()
  
 
class SyllabusCourseOutlineViewSet(viewsets.ModelViewSet):
    queryset = SyllabusCourseOutline.objects.all().order_by("row_no", "id")
    serializer_class = SyllabusCourseOutlineSerializer
    permission_classes = [RolePermission("ADMIN", "BAYANIHAN_LEADER")]

    def get_queryset(self):
        qs = super().get_queryset()
        syllabus_id = self.request.query_params.get("syllabus_id")
        syllabus_term = self.request.query_params.get("syllabus_term") 
        
        if syllabus_id:
            qs = qs.filter(syllabus_id=syllabus_id)
        if syllabus_term:
            qs = qs.filter(syllabus_term=syllabus_term.upper())   # normalize "midterm" → "MIDTERM"
            
        return qs
    
    # for Changing Order in the Course Outlines Table
    @action(detail=False, methods=["post"], url_path="reorder")
    def reorder(self, request):
        """
        Reorder course outlines for a syllabus.
        Expected payload:
        {
          "order": [
            {"id": 5, "position": 1},
            {"id": 8, "position": 2},
            {"id": 3, "position": 3}
          ]
        }
        """
        order = request.data.get("order", [])
        if not order:
            return Response({"detail": "No order provided."}, status=status.HTTP_400_BAD_REQUEST)

        outline_map = {item["id"]: item["position"] for item in order}

        outlines = SyllabusCourseOutline.objects.filter(id__in=outline_map.keys())
        for outline in outlines:
            outline.row_no = outline_map[outline.pk]
            outline.save(update_fields=["row_no"])

        return Response({"detail": "Order updated successfully."}, status=status.HTTP_200_OK)
    

class SyllabusTemplateViewSet(viewsets.ModelViewSet):
    queryset = SyllabusTemplate.objects.all().order_by("-created_at")
    serializer_class = SyllabusTemplateSerializer
    permission_classes = [RolePermission("ADMIN")]

    @action(detail=False, methods=["get"], url_path="latest")
    def latest_template(self, request):
        """Fetch the latest active template"""
        latest = SyllabusTemplate.objects.filter(is_active=True).order_by("-revision_no").first()
        if not latest:
            return Response({"detail": "No active syllabus template found."}, status=404)
        serializer = self.get_serializer(latest)
        return Response(serializer.data)

    def create(self, request, *args, **kwargs):
        """
        Treat 'create' as an update event — always create a new revision_no
        instead of modifying existing.
        """
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        self.perform_create(serializer)
        return Response(serializer.data, status=status.HTTP_201_CREATED)
    
    
class SyllabusCommentViewSet(viewsets.ModelViewSet):
    queryset = SyllabusComment.objects.select_related("user", "syllabus")
    serializer_class = SyllabusCommentSerializer
    permission_classes = [RolePermission("ADMIN", "BAYANIHAN_LEADER", "BAYANIHAN_TEACHER")]

    def perform_create(self, serializer):
        serializer.save(user=self.request.user)

    # ✅ Get all comments for a syllabus, optionally grouped by section
    @action(detail=False, methods=["get"])
    def by_syllabus(self, request):
        syllabus_id = request.query_params.get("syllabus_id")

        if not syllabus_id:
            return Response({"detail": "syllabus_id parameter is required."}, status=400)

        # Fetch all comments for that syllabus
        comments = self.queryset.filter(syllabus_id=syllabus_id)
        serializer = self.get_serializer(comments, many=True)

        # Optional: group by section
        grouped = {}
        for comment in serializer.data:
            section = comment["section"]
            grouped.setdefault(section, []).append(comment)

        return Response(grouped)

    @action(detail=False, methods=["get"])
    def by_section(self, request):
        syllabus_id = request.query_params.get("syllabus_id")
        section = request.query_params.get("section")

        if not syllabus_id or not section:
            return Response(
                {"detail": "syllabus_id and section parameters are required."},
                status=400,
            )

        comments = self.queryset.filter(syllabus_id=syllabus_id, section=section)
        serializer = self.get_serializer(comments, many=True)
        return Response(serializer.data)
    
    
class ReviewFormTemplateViewSet(viewsets.ModelViewSet):
    queryset = ReviewFormTemplate.objects.all().prefetch_related("items", "fields")
    serializer_class = ReviewFormTemplateSerializer
    permission_classes = [RolePermission("ADMIN", "CHAIRPERSON")]
 
    @action(detail=False, methods=["get"], url_path="latest-active")
    def latest_active(self, request):
        """Return the most recent active template (highest revision_no and is_active=True)."""
        latest = (
            ReviewFormTemplate.objects.filter(is_active=True)
            .order_by("-revision_no")
            .prefetch_related("items", "fields")
            .first()
        )
        if not latest:
            return Response({"detail": "No active review form template found."}, status=status.HTTP_404_NOT_FOUND)
        serializer = self.get_serializer(latest)
        return Response(serializer.data)

    @action(detail=True, methods=["post"])
    def clone(self, request, pk=None):
        """Duplicate an existing template to create a new revision."""
        original = self.get_object()
        new_revision_no = original.revision_no + 1
        clone = ReviewFormTemplate.objects.create(
            title=original.title,
            revision_no=new_revision_no,
            description=f"Cloned from revision_no {original.revision_no}",
            effective_date=None,
            is_active=False,
        )
        for item in original.items.all():
            ReviewFormItem.objects.create(
                form_template=clone,
                type=item.type,
                text=item.text,
                order=item.order,
            )
        for field in original.fields.all():
            ReviewFormField.objects.create(
                form_template=clone,
                label=field.label,
                field_type=field.field_type,
                is_required=field.is_required,
                position=field.position,
                display_order=field.display_order,
                row=field.row,
                column=field.column,
                span_full=field.spann_full, 
            )
        serializer = self.get_serializer(clone)
        return Response(serializer.data, status=status.HTTP_201_CREATED)


class SRFFormViewSet(viewsets.ModelViewSet):
    queryset = SRFForm.objects.all().select_related("form_template", "syllabus").prefetch_related("indicators", "field_values")
    serializer_class = SRFFormSerializer
    permission_classes = [RolePermission("ADMIN", "BAYANIHAN_TEACHER", "BAYANIHAN_LEADER", "CHAIRPERSON", "DEAN", "AUDITOR")]

    def get_queryset(self):
        syllabus_id = self.request.GET.get("syllabus_id")
        if syllabus_id:
            return SRFForm.objects.filter(syllabus_id=syllabus_id).select_related(
                "user", "syllabus"
            )
        return SRFForm.objects.all()

    # ✅ Custom action to fetch an SRFForm by syllabus ID
    @action(detail=False, methods=["get"], url_path="by-syllabus/(?P<syllabus_id>[^/.]+)")
    def get_by_syllabus(self, request, syllabus_id=None):
        """
        Fetch the SRFForm (review submission) related to a specific syllabus.
        Example: GET /api/srf-forms/by-syllabus/42/
        """
        srf_form = (
            SRFForm.objects.select_related("form_template", "syllabus", "user")
            .prefetch_related("indicators", "field_values")
            .filter(syllabus_id=syllabus_id)
            .first()  # ✅ This prevents DoesNotExist errors
        )

        if not srf_form:
            # Option 1: return empty response with 200
            return Response(None, status=status.HTTP_200_OK)
            # Option 2 (alternative): return 204 No Content
            # return Response(status=status.HTTP_204_NO_CONTENT)

        serializer = self.get_serializer(srf_form)
        return Response(serializer.data, status=status.HTTP_200_OK)
    
    @action(detail=True, methods=["get"])
    def export_docx(self, request, pk=None):
        # Windows-only COM initialization
        is_windows = platform.system().lower() == "windows"

        if is_windows and pythoncom is not None:
            pythoncom.CoInitialize()  
        
        try:
            review_form = self.get_object() 
            
            # 1️⃣ Load DOCX template
            template_path = os.path.join(
                settings.BASE_DIR, "syllabi", "templates", "ReviewFormHeaderTemp.docx"
            )
            doc = DocxTemplate(template_path)

            # 2️⃣ Fill in context in DOCX (simple replacement for placeholders)
            revision_no = (
                f"{int(review_form.form_template.revision_no):02d}"
                if review_form.form_template and review_form.form_template.revision_no is not None
                else ""
            )
            effective_date = (
                review_form.form_template.effective_date.strftime("%m.%d.%y")
                if review_form.form_template and review_form.form_template.effective_date
                else ""
            )  
            
            # Fill in your context 
            context = {
                "revision_no": revision_no,
                "effective_date": effective_date,
            } 
            doc.render(context) 
            
            # 7️⃣ Save to temp DOCX/PDF
            timestamp = review_form.review_date.strftime("%Y-%m-%d") if review_form.review_date else "NA"
            action_str = "APPROVED" if review_form.action == SRFForm.Action.APPROVED else "REJECTED"
            filename_docx = f"ReviewForm-{review_form.syllabus.course.course_code}-{action_str}-{timestamp}.docx"
            filename_pdf = filename_docx.replace(".docx", ".pdf") 

            # Use a temp directory so we can run LibreOffice/win32com conversion
            with tempfile.TemporaryDirectory() as tmpdir:
                temp_docx = os.path.join(tmpdir, filename_docx)
                temp_pdf = os.path.join(tmpdir, filename_pdf)

                # Save generated .docx locally for conversion
                doc.save(temp_docx)

                # Convert DOCX -> PDF
                convert_docx_to_pdf(temp_docx, temp_pdf)

                # Read PDF file
                with open(temp_pdf, "rb") as f:
                    pdf_bytes = f.read()

                # Upload to DigitalOcean Spaces
                pdf_path = default_storage.save(
                    f"review_forms/{filename_pdf}",  # folder inside bucket
                    ContentFile(pdf_bytes)
                )

            # Get the public URL for the browser to download
            pdf_url = default_storage.url(pdf_path)
 
            # 9️⃣ Return PDF as HTTP response
            return JsonResponse({"pdf_url": pdf_url})
            
        finally:
            if is_windows and pythoncom is not None:
                pythoncom.CoUninitialize()
